"""DOCX ingestion."""

from __future__ import annotations

import hashlib
import logging
from pathlib import Path

from docx import Document as DocxDocument

from docstack.models import DocumentRecord

logger = logging.getLogger(__name__)


def _doc_id(path: Path) -> str:
    st = path.stat()
    h = hashlib.sha256(f"{path.resolve()}:{st.st_mtime_ns}".encode()).hexdigest()
    return h[:16]


def extract_docx_records(path: Path) -> list[DocumentRecord]:
    doc_id = _doc_id(path)
    filename = path.name
    doc = DocxDocument(str(path))
    records: list[DocumentRecord] = []
    page = 1

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        block_type = "heading" if style.startswith("Heading") else "text"
        records.append(
            DocumentRecord(
                doc_id=doc_id,
                source_path=str(path.resolve()),
                filename=filename,
                mime="docx",
                page=page,
                block_type=block_type,
                text=text,
                section_heading=text if block_type == "heading" else None,
            )
        )

    for ti, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows.append("\t".join(cells))
        ttext = "\n".join(rows).strip()
        if ttext:
            records.append(
                DocumentRecord(
                    doc_id=doc_id,
                    source_path=str(path.resolve()),
                    filename=filename,
                    mime="docx",
                    page=page,
                    block_type="table",
                    text=ttext,
                    table_id=f"t{ti}",
                )
            )

    return records
